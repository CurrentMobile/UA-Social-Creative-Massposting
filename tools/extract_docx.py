"""Extract plain text / light markdown from a .docx file.

Step 1 (Input & Briefing) helper — turns a dropped Word brief into readable text the
briefing workflow can distill (see workflows/input_briefing.md). A .docx is just a ZIP of
XML, so this needs no third-party deps: it reads word/document.xml with the stdlib. If
python-docx happens to be installed it is used instead (slightly cleaner table handling).

Paragraphs become lines; Word tables become pipe-separated rows (one line per row) so the
structure survives. Headings/bullets are kept as plain text.

Usage:
    .venv\\Scripts\\python.exe tools/extract_docx.py "<file.docx>" [more.docx ...]
    .venv\\Scripts\\python.exe tools/extract_docx.py "<file.docx>" --out extracted.md

With no --out the text is printed to stdout. With multiple inputs each file is printed
under a "===== <name> =====" header (and --out is ignored — print only).
"""

from __future__ import annotations

import argparse
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _para_text(p: ET.Element) -> str:
    """Concatenate the text of one <w:p>, honouring tabs and line breaks."""
    out: list[str] = []
    for node in p.iter():
        tag = node.tag
        if tag == W + "t":
            out.append(node.text or "")
        elif tag == W + "tab":
            out.append("\t")
        elif tag in (W + "br", W + "cr"):
            out.append("\n")
    return "".join(out).strip()


def _table_text(tbl: ET.Element) -> list[str]:
    """Render a <w:tbl> as one pipe-separated line per row."""
    rows: list[str] = []
    for tr in tbl.findall(W + "tr"):
        cells = []
        for tc in tr.findall(W + "tc"):
            cell = " ".join(_para_text(p) for p in tc.findall(W + "p")).strip()
            cells.append(cell)
        if any(cells):
            rows.append(" | ".join(cells))
    return rows


def _extract_stdlib(path: Path) -> str:
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    body = root.find(W + "body")
    if body is None:
        return ""
    lines: list[str] = []
    for child in body:
        if child.tag == W + "p":
            lines.append(_para_text(child))
        elif child.tag == W + "tbl":
            lines.extend(_table_text(child))
    # collapse runs of blank lines to a single blank line
    cleaned: list[str] = []
    for ln in lines:
        if ln.strip() == "" and (not cleaned or cleaned[-1] == ""):
            continue
        cleaned.append(ln.rstrip())
    return "\n".join(cleaned).strip() + "\n"


def _extract_python_docx(path: Path) -> str:
    import docx  # type: ignore

    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    cleaned: list[str] = []
    for ln in lines:
        if ln.strip() == "" and (not cleaned or cleaned[-1] == ""):
            continue
        cleaned.append(ln.rstrip())
    return "\n".join(cleaned).strip() + "\n"


def extract(path: Path) -> str:
    """Return the document text, preferring python-docx, falling back to stdlib."""
    if not path.exists():
        raise FileNotFoundError(path)
    if path.suffix.lower() != ".docx":
        raise ValueError(f"not a .docx file: {path}")
    try:
        return _extract_python_docx(path)
    except ImportError:
        return _extract_stdlib(path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract text from .docx file(s)")
    ap.add_argument("files", type=Path, nargs="+", help=".docx file(s)")
    ap.add_argument("--out", type=Path, help="write text here (single input only)")
    args = ap.parse_args()

    if len(args.files) == 1 and args.out:
        text = extract(args.files[0])
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text, encoding="utf-8")
        print(f"wrote {len(text)} chars -> {args.out}")
        return 0

    for i, f in enumerate(args.files):
        if len(args.files) > 1:
            if i:
                print()
            print(f"===== {f.name} =====")
        print(extract(f), end="")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
